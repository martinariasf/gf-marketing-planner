# -*- coding: utf-8 -*-
"""GF on-brand .docx generator scaffold.

Reproduce with modifications: keep the helpers + GF palette, replace only the
CONTENIDO section at the bottom. Run with /opt/hermes/.venv/bin/python3.
Output goes under /opt/data/gf_docs/ (writable; /opt/hermes is not).
Install dep first:  cd /opt/hermes && VIRTUAL_ENV=/opt/hermes/.venv uv pip install python-docx -q
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- GF palette (brief branding.colors) ---
CHARCOAL = RGBColor(0x1a, 0x1a, 0x1a)
GREEN = RGBColor(0x22, 0xc5, 0x5e)
SLATE = RGBColor(0x64, 0x74, 0x8b)
SLATE_LIGHT = RGBColor(0x94, 0xa3, 0xb8)
WHITE = RGBColor(0xff, 0xff, 0xff)


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def green_rule(doc, sz='18'):
    pl = doc.add_paragraph()
    pPr = pl._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '22c55e')
    pbdr.append(bottom); pPr.append(pbdr)
    pl.paragraph_format.space_after = Pt(6)


def base_doc():
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Calibri'; st.font.size = Pt(11); st.font.color.rgb = CHARCOAL
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = Pt(54)
    sec.left_margin = sec.right_margin = Pt(54)
    return doc


def cover(doc, kicker, title, subtitle, meta):
    p = doc.add_paragraph(); r = p.add_run(kicker.upper())
    r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = GREEN
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph(); r = p.add_run(title)
    r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = CHARCOAL
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph(); r = p.add_run(subtitle)
    r.font.size = Pt(13); r.font.color.rgb = SLATE
    p.paragraph_format.space_after = Pt(6)
    green_rule(doc)
    p = doc.add_paragraph(); r = p.add_run(meta)
    r.font.size = Pt(9); r.font.color.rgb = SLATE_LIGHT
    doc.add_paragraph()


def h1(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = CHARCOAL
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(3)
    green_rule(doc, '6')


def h2(doc, text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(12.5); r.font.bold = True; r.font.color.rgb = GREEN
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)


def para(doc, text, size=11, color=CHARCOAL, bold=False, after=6, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold; r.font.color.rgb = color; r.font.italic = italic
    p.paragraph_format.space_after = Pt(after)
    return p


def bullet(doc, lead, text=''):
    p = doc.add_paragraph(style='List Bullet')
    if lead:
        r = p.add_run(lead); r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = CHARCOAL
    if text:
        r2 = p.add_run(text); r2.font.size = Pt(11); r2.font.color.rgb = CHARCOAL
    return p


def numbered(doc, lead, text=''):
    p = doc.add_paragraph(style='List Number')
    if lead:
        r = p.add_run(lead); r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = CHARCOAL
    if text:
        r2 = p.add_run(text); r2.font.size = Pt(11); r2.font.color.rgb = CHARCOAL
    return p


def callout(doc, title, text):
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.rows[0].cells[0]; shade(cell, '1a1a1a')
    cp = cell.paragraphs[0]; cp.text = ''
    r = cp.add_run(title); r.font.color.rgb = GREEN; r.font.size = Pt(11); r.font.bold = True
    cp2 = cell.add_paragraph(); r2 = cp2.add_run(text); r2.font.color.rgb = WHITE; r2.font.size = Pt(11.5)
    doc.add_paragraph()


# ==================== CONTENIDO (replace this block per document) ====================
doc = base_doc()
cover(doc,
      "Kicker en verde",
      "Titulo del documento",
      "Subtitulo descriptivo en slate",
      "GF Innovative Solutions  ·  Viktor  ·  Linea LATAM  ·  Documento interno")

h1(doc, "1. Primera seccion")
para(doc, "Texto de cuerpo. Voz de marca: frases de largo variado, sin em dash, sin buzzwords.")
h2(doc, "Subseccion")
numbered(doc, "Lead en negrita. ", "resto del punto.")
callout(doc, "La idea que debe aterrizar", "Texto blanco sobre caja charcoal.")

out = "/opt/data/gf_docs/GF_Documento.docx"
doc.save(out)
print("SAVED", out)
